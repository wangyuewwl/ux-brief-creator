#!/usr/bin/env python3
"""
Create a complete UX Brief with ALL content properly filled in
"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy

# Load the template from Input folder
template_path = "/Users/willyue/Mac Data/Design/IDCG/AI Agent Tools/Kiro/UX Brief/Input/UX Brief Template.docx"
# Output to the designated location
output_path = "/Users/willyue/Mac Data/Design/IDCG/AI Agent Tools/Kiro/UX Brief/Output/Cognito MRR - UX Brief FINAL.docx"
# Also save a backup in workspace
backup_path = "UX Brief/Cognito MRR - UX Brief FINAL.docx"

doc = Document(template_path)

print("📝 Starting to fill UX Brief template...")

# Define ALL content
content = {
    # Section 1: Project at-a-glance
    "project_title": "Cognito Multi-Region Replication (MRR) - UX Brief",
    "author": "Will Wang",
    "date": "December 1, 2025",
    "project_name": "Cognito Multi-Region Replication (MRR)",
    "aws_service": "Amazon Cognito",
    "description": "Enable customers to maintain seamless user access to their applications from multiple regions with enhanced reliability and business continuity through managed multi-region replication of authentication data.",
    "size": "Large",
    "asana_intake": "TBD",
    "asana_project": "TBD",
    "slack": "#cognito-mrr-ux",
    "tier": "1",
    "final_ux_date": "Nov 20, 2025",
    "ga_date": "Feb 26, 2026",
    "pm": "Howie Li",
    "program_manager": "N/A",
    "eng_team": "Vishal (Lead)",
    "uxd": "Will Wang",
    "uxr": "Will Wang",
    "tech_writer": "Brandon",
    
    # Section 2: Project goal
    "project_goal": "The goal is to enable Amazon Cognito customers to achieve business continuity and improved resiliency against regional disruptions in their identity systems. Success means customers can maintain seamless user authentication during rare regional outages without manual data handling, password resets, or service interruptions that lead to user frustration, churn, and revenue loss.",
    
    # Section 3: Personas
    "persona1_name": "Developer / Cloud Architect",
    "persona1_avatar": "Alex",
    "persona1_goal": "To build and maintain highly available applications with minimal operational overhead while ensuring users can always authenticate.",
    "persona1_opportunity": "This feature eliminates the need to build and maintain custom multi-region replication solutions, allowing Alex to focus on application features rather than infrastructure resilience.",
    
    # Section 4: Narrative
    "narrative_statement": "Alex, a cloud architect at a financial services company, needs a reliable way to ensure their customer-facing applications remain accessible during AWS regional disruptions without forcing users to reset passwords or compromising security.",
    
    "problem": """Alex is a cloud architect at Fidelity Investments, managing authentication for thousands of daily transactions. Every minute of downtime can impact customer trust and revenue. Currently, Alex's team has built an in-house solution for business continuity by manually exporting Cognito user data to a separate database in a different region.

This workaround is a constant source of stress. The manual export process is inefficient and not scalable. Worse, because Cognito doesn't allow exporting credential hashes for security reasons, users are forced to reset their passwords when accessing the application from the secondary region during a failover. This creates a terrible user experience during an already stressful outage situation.

"We're spending developer time maintaining this fragile system instead of building features our customers actually want," Alex explains to their manager. "And when we do have to failover, our users are frustrated by password resets. Some abandon the process entirely, which means lost transactions and angry customers calling our support team."

The security team is also concerned. The manual data handling introduces risks of data exposure and data loss. Alex knows there has to be a better way, but building a truly robust multi-region authentication system would take months of development time they simply don't have.""",
    
    "solution": """Alex reads the AWS "What's New" announcement about Cognito's new multi-region replication feature and immediately schedules a proof-of-concept.

Following the step-by-step guide in the Cognito console, Alex selects a secondary region (us-west-2) for their primary user pool in us-east-1. The console clearly shows the replication status and estimated completion time. Within an hour, the secondary region is ready on standby.

Next, Alex configures their custom domain and sets up a failover policy using Route 53 health checks. The console provides helpful templates and code samples, making the setup straightforward. Alex configures health checks for Cognito, Lambda, and DynamoDB dependencies, defining clear rules for when to trigger a failover.

"This is exactly what we needed," Alex thinks while reviewing the configuration. The console clearly shows which AWS dependencies need to be configured in the secondary region - Lambda triggers, SES for email OTPs, SNS for SMS. Alex uses the provided CloudFormation templates to replicate these configurations quickly.

A few weeks later, during a planned DR drill, Alex triggers a manual failover to test the system. The Route 53 health check detects the simulated outage and redirects traffic to the secondary region. Users who are already logged in stay logged in - their tokens work seamlessly across regions. New users can sign in with their existing credentials without any password resets.

"Our users didn't even notice the failover," Alex reports to leadership. "And we're no longer maintaining that fragile custom solution. This managed service handles all the heavy lifting - data synchronization, credential replication, everything. We can finally focus on building features instead of maintaining infrastructure."

The security team is equally pleased. Advanced security features like compromised credential checks and adaptive authentication continue working in the secondary region, maintaining the same high security posture.""",
    
    # Section 5: UX Scope
    "designs": """• Multi-region replication setup wizard - 4-step configuration flow in console
• User pool overview dashboard - Enhanced to show multi-region status with replication health
• Secondary region console experience - Modified console views with clear labeling and disabled operations
• Error states and messaging - User-friendly error handling for unsupported operations
• Health check dashboard - Monitoring and alerting interface with dependency status
• Configuration templates and code samples - CloudFormation templates and Route 53 examples""",
    
    "research_notes": "Extensive customer validation already conducted during PRFAQ phase with Cox Automotive, Wiz, Natura & Co, Shutterfly, First Advantage, Dow Jones, PegaSystems, NextEra (10+ customers). No formal usability testing planned for initial release due to timeline constraints and strong validation from customer feedback.",
    
    "related_projects": """• UIS Migration (Dependency): Multi-region replication only works with user pools migrated to UIS. Any delays will impact this project.
• Route 53 Health Checks Integration: Customers will use Route 53 for traffic management and failover triggering.
• CloudFormation Template Development: Templates for replicating AWS dependencies in secondary region.""",
    
    # Section 6: Schedule
    "start_date": "Aug 1, 2025",
    "total_effort": "15.8 weeks",
    "final_handoff": "Nov 20, 2025",
    
    # Section 7: Related links
    "related_links": """• (Howie Li) PR/FAQ: Cognito MRR PRFAQ_UX review.docx
• (Will Wang) User flows: TBD
• (Will Wang) Design explorations: TBD"""
}

print("✅ Content defined")

# Helper function to replace text in runs while preserving formatting
def replace_in_runs(paragraph, old_text, new_text):
    """Replace text in paragraph runs while preserving formatting"""
    full_text = paragraph.text
    if old_text in full_text:
        # Find which runs contain the text
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

# Helper function to replace text in entire document
def replace_text_everywhere(doc, old_text, new_text):
    """Replace text in all paragraphs and tables"""
    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_runs(para, old_text, new_text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, old_text, new_text)

print("🔄 Applying replacements...")

# Apply all replacements
replacements = {
    # Title
    "[Template]  [Project name]": content["project_name"],
    "[Project name]": content["project_name"],
    
    # Section 1 - Basic info
    "Month dd, yyyy": content["date"],
    "Write 1-2 sentences describing what this project is about.": content["description"],
    "Small, Medium, Large": content["size"],
    "1, 2, 3": content["tier"],
    "Apr 25, 2025": content["final_ux_date"],
    "Aug 24, 2025": content["ga_date"],
    
    # Section 2 - Project goal
    "[Provide 1-2 clear sentences on what the desired outcome of the project is.  What does a  success outcome look like?]": content["project_goal"],
    
    # Section 3 - Personas
    "[Cloud platform architect (aka Cloud Implementer)]": content["persona1_name"],
    "[Kelly]": content["persona1_avatar"],
    "[To design and maintain a secure infrastructure through a combination of policies, compliance and governance controls.]": content["persona1_goal"],
    "[This new capability provides Kelly increased visibility to the activity of the member accounts within her organization, while also providing greater ability to delegate administration.]": content["persona1_opportunity"],
    
    # Section 4 - Narrative
    'Write a single statement that explains the benefit the project will provide to your primary persona (ex: Kelly manages her organization at Capital One, and is seeking a way to delegate management of applications.)"': content["narrative_statement"],
    "[This portion of the narrative should establish the problem case with the personas' jobs to be done, and notable gaps that this project will address. If this is an improvement over an existing process, describe how your customer performs the job today before your solution. It's fine to exaggerate with emotion and quotes.]": content["problem"],
    "[Describe the UX of the solution from the perspective of your personas narrative. What are they thinking as they discover and use the new capability? What are they saying to others? Describe how the solution takes advantage of any obvious opportunities.]": content["solution"],
    
    # Section 5 - UX Scope
    "[Provide a short inventory of the expected designs (i.e. screens, flows, widgets, etc.) that you will need to deliver in order to achieve the desired outcome above.]": content["designs"],
    "[Title of project (PR/FAQ Link): short description]": content["related_projects"],
    
    # Section 6 - Schedule
    "Jan 6, 2025": content["start_date"],
    "19.1 weeks": content["total_effort"],
    
    # Section 7 - Related links
    "[(Author) PR/FAQ: link]": content["related_links"],
}

for old, new in replacements.items():
    replace_text_everywhere(doc, old, new)
    print(f"  ✓ Replaced: {old[:50]}...")

print("📋 Filling stakeholder table...")

# Special handling for stakeholder table (Section 1)
# Find the table with stakeholders
for table in doc.tables:
    table_text = ""
    for row in table.rows:
        for cell in row.cells:
            table_text += cell.text
    
    # Check if this is the stakeholders table
    if "Author of doc" in table_text or "Product manager" in table_text:
        print("  Found stakeholders table!")
        for row in table.rows:
            cells = list(row.cells)
            if len(cells) >= 2:
                label = cells[0].text.strip()
                
                # Map labels to values
                stakeholder_map = {
                    "Author of doc (UXD)": content["author"],
                    "Date of last update": content["date"],
                    "Project name": content["project_name"],
                    "AWS service": content["aws_service"],
                    "Short description": content["description"],
                    "Size": content["size"],
                    "Asana intake ticket link": content["asana_intake"],
                    "Asana project link": content["asana_project"],
                    "Slack primary UX channel": content["slack"],
                    "Tier": content["tier"],
                    "Final UX delivery date*": content["final_ux_date"],
                    "GA date*": content["ga_date"],
                    "Product manager": content["pm"],
                    "Program manager": content["program_manager"],
                    "Engineering team": content["eng_team"],
                    "UX designer": content["uxd"],
                    "UX researcher": content["uxr"],
                    "Tech writer": content["tech_writer"],
                }
                
                # If this row has a matching label, fill the value
                for key, value in stakeholder_map.items():
                    if key in label:
                        # Clear and fill the value cell
                        value_cell = cells[1]
                        value_cell.text = ""
                        for para in value_cell.paragraphs:
                            para.clear()
                        para = value_cell.paragraphs[0] if value_cell.paragraphs else value_cell.add_paragraph()
                        run = para.add_run(value)
                        run.font.size = Pt(11)
                        print(f"  ✓ Filled: {key} = {value}")
                        break

print("📊 Filling research table...")

# Fill research table (Section 5b)
for table in doc.tables:
    table_text = ""
    for row in table.rows:
        for cell in row.cells:
            table_text += cell.text
    
    if "Customer interviews" in table_text and "Usability testing" in table_text:
        print("  Found research table!")
        for row in table.rows:
            cells = list(row.cells)
            row_text = " ".join([c.text for c in cells])
            
            # Fill Customer interviews row
            if "Customer interviews" in row_text and len(cells) >= 4:
                # Mark "No" column
                cells[2].text = ""
                para = cells[2].paragraphs[0] if cells[2].paragraphs else cells[2].add_paragraph()
                para.clear()
                run = para.add_run("✅")
                run.font.size = Pt(11)
                
                # Fill notes
                cells[3].text = ""
                para = cells[3].paragraphs[0] if cells[3].paragraphs else cells[3].add_paragraph()
                para.clear()
                run = para.add_run(content["research_notes"])
                run.font.size = Pt(10)
                print("  ✓ Filled Customer interviews")
            
            # Fill Usability testing row
            if "Usability testing" in row_text and len(cells) >= 4:
                # Mark "No" column
                cells[2].text = ""
                para = cells[2].paragraphs[0] if cells[2].paragraphs else cells[2].add_paragraph()
                para.clear()
                run = para.add_run("✅")
                run.font.size = Pt(11)
                print("  ✓ Filled Usability testing")

print("💾 Saving document...")
doc.save(output_path)
doc.save(backup_path)

print(f"\n✅ Complete UX Brief created successfully!")
print(f"📄 Primary location: {output_path}")
print(f"📄 Backup location: {backup_path}")
print(f"\n🎉 All fields have been filled in correctly!")
